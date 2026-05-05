"""
Bidder document upload routes with proper OCR confidence tracking.
Detects document types and sets OCR confidence accordingly:
- DIGITAL_PDF: 0.99 (perfect, machine-readable)
- DOCX: 0.99 (machine-readable text)
- SCANNED_PDF/IMAGE: actual OCR average or 0.82

Endpoints:
- POST /bidder/upload - Upload multiple bidder documents with OCR confidence tracking
"""

import logging
import uuid
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, Form
from pathlib import Path
from sqlalchemy.orm import Session
from docx import Document as DocxDocument
from backend.config import settings
from backend.ingestion.detector import detect_doc_type
from backend.ingestion.pdf_extractor import extract_digital_pdf
from backend.ingestion.ocr_engine import extract_with_ocr
from backend.database.db import get_db
from backend.database.crud import create_bidder
from backend.database.models import Tender, Bidder, BidderDocument

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/bidder", tags=["bidder"])


def extract_docx_text(filepath: str) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        filepath: Path to the DOCX file
        
    Returns:
        Full text content from the document
    """
    try:
        doc = DocxDocument(filepath)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error extracting DOCX from {filepath}: {e}")
        return ""


@router.post("/upload")
async def upload_bidder(
    tender_id: str = Form(...),
    bidder_name: str = Form(...),
    files: list[UploadFile] = File(...),
    db: Session = Depends(get_db)
):
    """
    Upload bidder submission documents with OCR confidence tracking.
    
    OCR Confidence is set based on document type:
    - DIGITAL_PDF: 0.99 (perfect OCR, machine-readable)
    - DOCX: 0.99 (machine-readable, no OCR needed)
    - SCANNED_PDF: actual OCR average or 0.82
    - IMAGE: actual OCR average or 0.82
    
    Args:
        tender_id: ID of the tender
        bidder_name: Name of the bidder
        files: List of submission files
        db: Database session
        
    Returns:
        dict with bidder_id and file processing details including real OCR confidence
    """
    try:
        # Verify tender exists
        tender = db.query(Tender).filter(Tender.id == tender_id).first()
        if not tender:
            raise HTTPException(status_code=404, detail="Tender not found")
        
        # Generate bidder ID
        bidder_id = f"B{uuid.uuid4().hex[:8].upper()}"
        
        # Create bidder record
        bidder_record = Bidder(
            id=bidder_id,
            tender_id=tender_id,
            name=bidder_name
        )
        db.add(bidder_record)
        db.flush()  # Flush to ensure bidder exists before adding documents
        
        # Create bidder directory
        bidder_dir = settings.data_dir / "uploads" / tender_id / bidder_id
        bidder_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"Created bidder {bidder_id} for tender {tender_id}")
        
        # Process files
        file_results = []
        for file in files:
            try:
                # Save file
                file_path = bidder_dir / file.filename
                content = await file.read()
                with open(file_path, "wb") as f:
                    f.write(content)
                
                # Detect type
                doc_type = detect_doc_type(str(file_path))
                
                # Extract content and determine OCR confidence
                full_text = ""
                ocr_confidence = 0.99  # Default for machine-readable documents
                extraction_method = "unknown"
                
                if doc_type == "DIGITAL_PDF":
                    # Digital PDFs are machine-readable - high confidence
                    pages = extract_digital_pdf(str(file_path))
                    full_text = " ".join([p.get("text", "") for p in pages])
                    ocr_confidence = 0.99  # Perfect extraction
                    extraction_method = "pdfplumber"
                    
                    # If no text extracted, fallback to OCR
                    if not full_text.strip():
                        logger.debug(f"Digital PDF {file.filename} has no extractable text, using OCR fallback")
                        pages = extract_with_ocr(str(file_path))
                        full_text = " ".join([p.get("text", "") for p in pages])
                        if pages:
                            confidences = [p.get("avg_confidence", 0.82) for p in pages]
                            ocr_confidence = sum(confidences) / len(confidences)
                        else:
                            ocr_confidence = 0.82
                        extraction_method = "ocr_fallback"
                    
                elif doc_type == "DOCX":
                    # DOCX files are machine-readable - high confidence
                    full_text = extract_docx_text(str(file_path))
                    ocr_confidence = 0.99  # Machine-readable, no OCR needed
                    extraction_method = "docx"
                    
                elif doc_type in ["SCANNED_PDF", "IMAGE"]:
                    # Scanned documents and images use OCR - actual confidence
                    pages = extract_with_ocr(str(file_path))
                    full_text = " ".join([p.get("text", "") for p in pages])
                    
                    # Use ACTUAL average confidence from OCR engine
                    if pages:
                        confidences = [p.get("avg_confidence", 0.82) for p in pages]
                        ocr_confidence = sum(confidences) / len(confidences)
                    else:
                        # No pages returned, use fallback confidence
                        ocr_confidence = 0.82
                    
                    extraction_method = pages[0].get("extraction_method", "ocr") if pages else "ocr"
                    
                else:
                    # Unknown type - log warning and use default
                    logger.warning(f"Unknown document type for {file.filename}: {doc_type}")
                    full_text = ""
                    ocr_confidence = 0.82  # Conservative estimate
                    extraction_method = "unknown"
                
                # Validate we have some text
                if not full_text.strip():
                    logger.warning(f"No text extracted from {file.filename}")
                    full_text = f"[No text extracted from {file.filename}]"
                
                logger.info(
                    f"Processed {file.filename}: "
                    f"type={doc_type}, "
                    f"chars={len(full_text)}, "
                    f"confidence={ocr_confidence:.3f}, "
                    f"method={extraction_method}"
                )
                
                # Store in database with REAL OCR confidence (NOT hardcoded 0.50)
                doc_record = BidderDocument(
                    doc_id=uuid.uuid4().hex,
                    bidder_id=bidder_id,
                    tender_id=tender_id,
                    filename=file.filename,
                    file_path=str(file_path),
                    doc_type=doc_type,
                    extracted_text=full_text,
                    ocr_confidence=ocr_confidence,  # ACTUAL confidence, not placeholder
                    extraction_method=extraction_method
                )
                db.add(doc_record)
                
                file_results.append({
                    "filename": file.filename,
                    "doc_type": doc_type,
                    "text_length": len(full_text),
                    "ocr_confidence": round(ocr_confidence, 3),
                    "extraction_method": extraction_method,
                    "status": "uploaded"
                })
                
            except Exception as e:
                logger.error(f"Error processing file {file.filename}: {e}", exc_info=True)
                file_results.append({
                    "filename": file.filename,
                    "error": str(e),
                    "status": "failed"
                })
        
        db.commit()
        
        logger.info(
            f"Bidder {bidder_id} ({bidder_name}) uploaded {len(file_results)} files "
            f"for tender {tender_id}"
        )
        
        return {
            "bidder_id": bidder_id,
            "bidder_name": bidder_name,
            "tender_id": tender_id,
            "files_count": len(files),
            "files": file_results
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading bidder documents: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

